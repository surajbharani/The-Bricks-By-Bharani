"""
Agent Nano Bricks — Tool Executor (Hermes-grade tool set)

~20 tools, all sandboxed to the workspace:
  Files:   read_file, write_file, edit_file, multi_edit, append_file,
           delete_file, move_file, copy_file, make_dir, list_dir
  Search:  search_text (grep), find_files (glob)
  Exec:    shell_exec, run_python
  Web:     web_fetch, web_search, browser_action, generate_image
  Agents:  spawn_subagent (delegate a focused subtask)
"""
import json
import os
import re
import shutil
import subprocess
import urllib.parse
import urllib.request
import urllib.error
from pathlib import Path
from typing import Any, Optional


class SandboxViolation(Exception):
    pass


def _safe_path(workspace: Path, rel_or_abs: str) -> Path:
    p = Path(rel_or_abs)
    if not p.is_absolute():
        p = workspace / p
    resolved = p.resolve()
    try:
        resolved.relative_to(workspace.resolve())
    except ValueError:
        raise SandboxViolation(f"Path '{rel_or_abs}' is outside the workspace.")
    return resolved


_SKIP_DIRS = {".git", "node_modules", "__pycache__", ".venv", "venv", "dist", "build",
              ".nanobricks_memory", ".nanobricks_checkpoints"}
_BINARY_EXT = {".png", ".jpg", ".jpeg", ".gif", ".pdf", ".zip", ".exe", ".bin", ".so", ".dll", ".woff", ".ico", ".mp4", ".mp3"}


# ── File Tools ────────────────────────────────────────────────────────────────

def read_file(workspace: Path, path: str, offset: int = 0, limit: int = 0) -> dict[str, Any]:
    try:
        target = _safe_path(workspace, path)
        content = target.read_text(encoding="utf-8", errors="replace")
        if offset or limit:
            lines = content.splitlines()
            lines = lines[offset:offset + limit] if limit else lines[offset:]
            content = "\n".join(lines)
        if len(content) > 16000:
            content = content[:16000] + f"\n...[truncated, {len(content)-16000} chars omitted]"
        return {"ok": True, "content": content, "path": str(target)}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except FileNotFoundError:
        return {"ok": False, "error": f"File not found: {path}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def write_file(workspace: Path, path: str, content: str) -> dict[str, Any]:
    try:
        target = _safe_path(workspace, path)
        target.parent.mkdir(parents=True, exist_ok=True)
        existed = target.exists()
        target.write_text(content, encoding="utf-8")
        return {"ok": True, "action": "edit" if existed else "write", "path": str(target)}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def edit_file(workspace: Path, path: str, old_string: str, new_string: str,
              replace_all: bool = False) -> dict[str, Any]:
    """Precise line-level edit: replace an exact substring. Fails if the target
    text is missing or ambiguous (appears >1 time and replace_all is False)."""
    try:
        target = _safe_path(workspace, path)
        if not target.exists():
            return {"ok": False, "error": f"File not found: {path}"}
        content = target.read_text(encoding="utf-8", errors="replace")
        count = content.count(old_string)
        if count == 0:
            return {"ok": False, "error": "old_string not found in file. Read the file and copy the exact text."}
        if count > 1 and not replace_all:
            return {"ok": False, "error": f"old_string appears {count} times — not unique. Add more context or set replace_all=true."}
        new_content = content.replace(old_string, new_string) if replace_all \
            else content.replace(old_string, new_string, 1)
        target.write_text(new_content, encoding="utf-8")
        return {"ok": True, "action": "edit", "path": str(target), "replacements": count if replace_all else 1}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def multi_edit(workspace: Path, path: str, edits: list) -> dict[str, Any]:
    """Apply several edits to one file atomically — all succeed or none are written."""
    try:
        target = _safe_path(workspace, path)
        if not target.exists():
            return {"ok": False, "error": f"File not found: {path}"}
        content = target.read_text(encoding="utf-8", errors="replace")
        applied = 0
        for i, e in enumerate(edits or []):
            old = e.get("old_string", "")
            new = e.get("new_string", "")
            ra = e.get("replace_all", False)
            if not old:
                return {"ok": False, "error": f"Edit {i+1}: old_string is empty."}
            cnt = content.count(old)
            if cnt == 0:
                return {"ok": False, "error": f"Edit {i+1}: old_string not found."}
            if cnt > 1 and not ra:
                return {"ok": False, "error": f"Edit {i+1}: old_string not unique ({cnt}×)."}
            content = content.replace(old, new) if ra else content.replace(old, new, 1)
            applied += 1
        target.write_text(content, encoding="utf-8")
        return {"ok": True, "action": "edit", "path": str(target), "edits_applied": applied}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def append_file(workspace: Path, path: str, content: str) -> dict[str, Any]:
    try:
        target = _safe_path(workspace, path)
        target.parent.mkdir(parents=True, exist_ok=True)
        with open(target, "a", encoding="utf-8") as f:
            f.write(content)
        return {"ok": True, "path": str(target)}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def delete_file(workspace: Path, path: str) -> dict[str, Any]:
    try:
        target = _safe_path(workspace, path)
        if target == workspace.resolve():
            return {"ok": False, "error": "Refusing to delete the workspace root. Delete specific files/folders instead."}
        if not target.exists():
            return {"ok": False, "error": f"Not found: {path}"}
        if target.is_dir():
            shutil.rmtree(target)
        else:
            target.unlink()
        return {"ok": True, "deleted": str(target)}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def move_file(workspace: Path, src: str, dst: str) -> dict[str, Any]:
    try:
        s = _safe_path(workspace, src)
        d = _safe_path(workspace, dst)
        if s == workspace.resolve():
            return {"ok": False, "error": "Refusing to move the workspace root."}
        if not s.exists():
            return {"ok": False, "error": f"Source not found: {src}"}
        d.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(s), str(d))
        return {"ok": True, "from": str(s), "to": str(d)}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def copy_file(workspace: Path, src: str, dst: str) -> dict[str, Any]:
    try:
        s = _safe_path(workspace, src)
        d = _safe_path(workspace, dst)
        if not s.exists():
            return {"ok": False, "error": f"Source not found: {src}"}
        d.parent.mkdir(parents=True, exist_ok=True)
        if s.is_dir():
            shutil.copytree(str(s), str(d), dirs_exist_ok=True)
        else:
            shutil.copy2(str(s), str(d))
        return {"ok": True, "from": str(s), "to": str(d)}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def make_dir(workspace: Path, path: str) -> dict[str, Any]:
    try:
        target = _safe_path(workspace, path)
        target.mkdir(parents=True, exist_ok=True)
        return {"ok": True, "path": str(target)}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def list_dir(workspace: Path, path: str = ".") -> dict[str, Any]:
    try:
        target = _safe_path(workspace, path)
        if not target.is_dir():
            return {"ok": False, "error": f"Not a directory: {path}"}
        entries = []
        for item in sorted(target.iterdir()):
            try:
                size = item.stat().st_size if item.is_file() else None
            except Exception:
                size = None
            entries.append({"name": item.name, "type": "dir" if item.is_dir() else "file", "size": size})
        return {"ok": True, "entries": entries[:200]}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ── Search Tools ──────────────────────────────────────────────────────────────

def search_text(workspace: Path, pattern: str, path: str = ".",
                glob: str = "", max_results: int = 100) -> dict[str, Any]:
    """Regex search file contents (grep). Returns matching lines with locations."""
    try:
        root = _safe_path(workspace, path)
        rx = re.compile(pattern)
        results = []
        for dirpath, dirnames, filenames in os.walk(root):
            dirnames[:] = [d for d in dirnames if d not in _SKIP_DIRS]
            for fn in filenames:
                if Path(fn).suffix.lower() in _BINARY_EXT:
                    continue
                if glob and not Path(fn).match(glob):
                    continue
                fp = Path(dirpath) / fn
                try:
                    with open(fp, "r", encoding="utf-8", errors="ignore") as f:
                        for ln, line in enumerate(f, 1):
                            if rx.search(line):
                                results.append({
                                    "file": str(fp.relative_to(workspace)),
                                    "line": ln,
                                    "text": line.rstrip()[:300],
                                })
                                if len(results) >= max_results:
                                    return {"ok": True, "matches": results, "truncated": True}
                except Exception:
                    continue
        return {"ok": True, "matches": results, "count": len(results)}
    except re.error as e:
        return {"ok": False, "error": f"Invalid regex: {e}"}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def find_files(workspace: Path, pattern: str, path: str = ".") -> dict[str, Any]:
    """Find files by glob pattern, e.g. '**/*.py' or 'src/*.ts'."""
    try:
        root = _safe_path(workspace, path)
        matches = []
        for p in root.glob(pattern):
            if any(part in _SKIP_DIRS for part in p.parts):
                continue
            matches.append(str(p.relative_to(workspace)))
            if len(matches) >= 200:
                break
        return {"ok": True, "files": sorted(matches), "count": len(matches)}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ── Exec Tools ────────────────────────────────────────────────────────────────

def shell_exec(workspace: Path, command: str, timeout: int = 120) -> dict[str, Any]:
    try:
        if "../" in command or "..\\" in command:
            return {"ok": False, "error": "Path traversal in command is not allowed."}
        # Make _nb_utils importable from shell Python scripts via PYTHONPATH
        nb_path = str(workspace / "_nb_utils")
        existing_pypath = os.environ.get("PYTHONPATH", "")
        pypath = nb_path + os.pathsep + existing_pypath if existing_pypath else nb_path
        result = subprocess.run(
            command, shell=True, cwd=str(workspace),
            capture_output=True, text=True, timeout=timeout,
            env={**os.environ, "PYTHONIOENCODING": "utf-8", "PYTHONPATH": pypath},
        )
        output = (result.stdout + result.stderr).strip()
        if len(output) > 12000:
            output = output[:12000] + f"\n...[truncated, {len(output)-12000} chars omitted]"
        return {"ok": result.returncode == 0, "returncode": result.returncode, "output": output}
    except subprocess.TimeoutExpired:
        return {"ok": False, "error": f"Command timed out after {timeout}s."}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def run_python(workspace: Path, code: str, timeout: int = 120) -> dict[str, Any]:
    """Run a Python snippet directly (sub-script). Output capped at 12K chars."""
    try:
        # Prepend _nb_utils to sys.path so `import csv_utils` etc. work out of the box
        nb_path = str(workspace / "_nb_utils")
        preamble = f"import sys; sys.path.insert(0, {nb_path!r})\n"
        result = subprocess.run(
            ["python3", "-c", preamble + code], cwd=str(workspace),
            capture_output=True, text=True, timeout=timeout,
            env={**os.environ, "PYTHONIOENCODING": "utf-8"},
        )
        output = (result.stdout + result.stderr).strip()
        if len(output) > 12000:
            output = output[:12000] + f"\n...[truncated]"
        return {"ok": result.returncode == 0, "returncode": result.returncode, "output": output}
    except subprocess.TimeoutExpired:
        return {"ok": False, "error": f"Python script timed out after {timeout}s."}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ── Web Tools ─────────────────────────────────────────────────────────────────

def web_fetch(workspace: Path, url: str, method: str = "GET",
              headers: dict | None = None, body: str | None = None) -> dict[str, Any]:
    try:
        req = urllib.request.Request(url, method=method.upper())
        req.add_header("User-Agent", "Mozilla/5.0 (Agent Nano Bricks/1.0; compatible)")
        req.add_header("Accept", "text/html,application/json,*/*;q=0.9")
        if headers:
            for k, v in headers.items():
                req.add_header(k, str(v))
        if body:
            req.data = body.encode("utf-8")
            if not headers or "Content-Type" not in headers:
                req.add_header("Content-Type", "application/json")
        with urllib.request.urlopen(req, timeout=30) as resp:
            raw = resp.read(81920)
            charset = "utf-8"
            ct = resp.headers.get("Content-Type", "")
            if "charset=" in ct:
                charset = ct.split("charset=")[-1].split(";")[0].strip()
            try:
                text = raw.decode(charset, errors="replace")
            except (LookupError, UnicodeDecodeError):
                text = raw.decode("utf-8", errors="replace")
            return {"ok": True, "status": resp.status, "url": resp.url, "content": text[:40000]}
    except urllib.error.HTTPError as e:
        body_bytes = b""
        try:
            body_bytes = e.read(4096)
        except Exception:
            pass
        return {"ok": False, "error": f"HTTP {e.code}: {e.reason}", "body": body_bytes.decode("utf-8", errors="replace")}
    except urllib.error.URLError as e:
        return {"ok": False, "error": f"URL error: {e.reason}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def web_search(workspace: Path, query: str, max_results: int = 8) -> dict[str, Any]:
    """Search the web (DuckDuckGo HTML endpoint — no API key)."""
    try:
        url = "https://html.duckduckgo.com/html/?q=" + urllib.parse.quote(query)
        req = urllib.request.Request(url, headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        })
        with urllib.request.urlopen(req, timeout=20) as resp:
            html = resp.read().decode("utf-8", errors="replace")
        results = []
        # Parse result anchors
        for m in re.finditer(r'<a[^>]*class="result__a"[^>]*href="([^"]+)"[^>]*>(.*?)</a>', html, re.DOTALL):
            href, title = m.group(1), re.sub(r"<[^>]+>", "", m.group(2)).strip()
            # DuckDuckGo wraps real URL in a redirect param
            real = urllib.parse.unquote(href)
            mm = re.search(r"uddg=([^&]+)", real)
            link = urllib.parse.unquote(mm.group(1)) if mm else real
            if title:
                results.append({"title": title[:200], "url": link})
            if len(results) >= max_results:
                break
        if not results:
            return {"ok": True, "results": [], "note": "No results parsed — try web_fetch on a known URL."}
        return {"ok": True, "results": results}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def generate_image(workspace: Path, prompt: str, path: str = "image.png",
                   width: int = 1024, height: int = 1024) -> dict[str, Any]:
    """Generate an image from a text prompt and save it to the workspace.
    Uses the keyless Pollinations image service."""
    try:
        target = _safe_path(workspace, path)
        target.parent.mkdir(parents=True, exist_ok=True)
        enc = urllib.parse.quote(prompt)
        url = f"https://image.pollinations.ai/prompt/{enc}?width={width}&height={height}&nologo=true"
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0 (Agent Nano Bricks)"})
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = resp.read()
        if not data or len(data) < 100:
            return {"ok": False, "error": "Image service returned no data."}
        target.write_bytes(data)
        return {"ok": True, "path": str(target), "bytes": len(data)}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": f"Image generation failed: {e}"}


# ── Browser Tool ──────────────────────────────────────────────────────────────

_BROWSER_SCRIPT = '''
import sys, json
try:
    from playwright.sync_api import sync_playwright
except ImportError:
    print(json.dumps({{"ok": False, "error": "Playwright not installed. Run: pip install playwright && playwright install chromium"}}))
    sys.exit(0)

action = {action}
url = {url}
selector = {selector}
value = {value}
script = {script}
save_path = {save_path}

with sync_playwright() as p:
    browser = p.chromium.launch(headless=True, args=["--no-sandbox", "--disable-dev-shm-usage"])
    ctx = browser.new_context(viewport={{"width": 1280, "height": 800}})
    page = ctx.new_page()
    result = {{"ok": True}}
    try:
        if action == "navigate":
            page.goto(url, wait_until="domcontentloaded", timeout=30000)
            page.wait_for_load_state("networkidle", timeout=10000)
            result["title"] = page.title(); result["url"] = page.url
        elif action == "click":
            page.locator(selector).first.click(timeout=15000); result["clicked"] = selector
        elif action == "fill":
            page.locator(selector).first.fill(value, timeout=15000); result["filled"] = selector
        elif action == "get_text":
            text = page.locator(selector).first.inner_text(timeout=10000) if selector else page.inner_text("body")
            result["text"] = text[:10000]
        elif action == "screenshot":
            page.screenshot(path=save_path, full_page=True); result["path"] = save_path
        elif action == "evaluate":
            result["value"] = str(page.evaluate(script))[:4000]
        elif action == "get_page":
            result["html"] = page.content()[:20000]
        else:
            result = {{"ok": False, "error": f"Unknown action: {{action}}"}}
    except Exception as ex:
        result = {{"ok": False, "error": str(ex)}}
    browser.close()
    print(json.dumps(result))
'''


def browser_action(workspace: Path, action: str, url: str = "", selector: str = "",
                   value: str = "", script: str = "", screenshot_path: str = "") -> dict[str, Any]:
    try:
        save_path = screenshot_path if screenshot_path else str(workspace / "screenshot.png")
        code = _BROWSER_SCRIPT.format(
            action=json.dumps(action), url=json.dumps(url), selector=json.dumps(selector),
            value=json.dumps(value), script=json.dumps(script), save_path=json.dumps(save_path),
        )
        proc = subprocess.run(["python3", "-c", code], capture_output=True, text=True,
                              timeout=90, cwd=str(workspace))
        out = proc.stdout.strip()
        if out:
            try:
                return json.loads(out)
            except json.JSONDecodeError:
                return {"ok": True, "output": out[:4000]}
        if proc.returncode != 0:
            err = proc.stderr.strip()
            return {"ok": False, "error": err[:2000] if err else "Browser process failed."}
        return {"ok": True, "output": "(no output)"}
    except subprocess.TimeoutExpired:
        return {"ok": False, "error": "Browser action timed out after 90s."}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ── Document Tools ────────────────────────────────────────────────────────────

def read_document(workspace: Path, path: str) -> dict[str, Any]:
    """Read PDF / Word / Excel / CSV / text and return its text content.
    Degrades with a clear, actionable error if an optional library is missing."""
    try:
        target = _safe_path(workspace, path)
        if not target.exists():
            return {"ok": False, "error": f"File not found: {path}"}
        ext = target.suffix.lower()

        if ext == ".pdf":
            try:
                try:
                    from pypdf import PdfReader
                except ImportError:
                    from PyPDF2 import PdfReader
            except ImportError:
                return {"ok": False, "error": "PDF support needs pypdf. Run: pip install pypdf"}
            except BaseException as e:
                return {"ok": False, "error": f"PDF library failed to load: {e}"}
            reader = PdfReader(str(target))
            pages = [(p.extract_text() or "") for p in reader.pages]
            text = "\n\n".join(pages)
            return {"ok": True, "content": text[:40000], "pages": len(pages)}

        if ext in (".docx",):
            try:
                import docx
            except ImportError:
                return {"ok": False, "error": "Word support needs python-docx. Run: pip install python-docx"}
            d = docx.Document(str(target))
            text = "\n".join(p.text for p in d.paragraphs)
            return {"ok": True, "content": text[:40000], "paragraphs": len(d.paragraphs)}

        if ext in (".xlsx", ".xlsm"):
            try:
                import openpyxl
            except ImportError:
                return {"ok": False, "error": "Excel support needs openpyxl. Run: pip install openpyxl"}
            wb = openpyxl.load_workbook(str(target), read_only=True, data_only=True)
            out = []
            for ws in wb.worksheets:
                out.append(f"# Sheet: {ws.title}")
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i >= 200:
                        out.append("...[more rows omitted]")
                        break
                    out.append("\t".join("" if c is None else str(c) for c in row))
            return {"ok": True, "content": "\n".join(out)[:40000], "sheets": len(wb.worksheets)}

        # csv / tsv / txt / md / json / etc → plain text
        text = target.read_text(encoding="utf-8", errors="replace")
        return {"ok": True, "content": text[:40000]}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": f"Could not read document: {e}"}


def generate_document(workspace: Path, path: str, content: str, title: str = "") -> dict[str, Any]:
    """Create a PDF or Word (.docx) document from text/markdown-ish content.
    Format chosen by the output file extension."""
    try:
        target = _safe_path(workspace, path)
        target.parent.mkdir(parents=True, exist_ok=True)
        ext = target.suffix.lower()

        if ext == ".pdf":
            try:
                from fpdf import FPDF
            except ImportError:
                return {"ok": False, "error": "PDF generation needs fpdf2. Run: pip install fpdf2"}
            except BaseException as e:
                return {"ok": False, "error": f"PDF library failed to load: {e}"}
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            if title:
                pdf.set_font("Helvetica", "B", 16)
                pdf.multi_cell(0, 10, title)
                pdf.ln(2)
            pdf.set_font("Helvetica", size=11)
            for line in content.split("\n"):
                safe = line.encode("latin-1", "replace").decode("latin-1")
                pdf.multi_cell(0, 6, safe if safe else " ")
            pdf.output(str(target))
            return {"ok": True, "path": str(target), "format": "pdf"}

        if ext == ".docx":
            try:
                import docx
            except ImportError:
                return {"ok": False, "error": "Word generation needs python-docx. Run: pip install python-docx"}
            d = docx.Document()
            if title:
                d.add_heading(title, level=0)
            for line in content.split("\n"):
                if line.startswith("# "):
                    d.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    d.add_heading(line[3:], level=2)
                else:
                    d.add_paragraph(line)
            d.save(str(target))
            return {"ok": True, "path": str(target), "format": "docx"}

        # Fallback: plain text / markdown
        if title:
            content = f"# {title}\n\n{content}"
        target.write_text(content, encoding="utf-8")
        return {"ok": True, "path": str(target), "format": "text"}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": f"Could not generate document: {e}"}


def analyze_data(workspace: Path, path: str, question: str = "") -> dict[str, Any]:
    """Analyze a CSV/Excel file: columns, row count, and numeric summaries.
    Uses pandas if available, else a stdlib fallback for CSV."""
    try:
        target = _safe_path(workspace, path)
        if not target.exists():
            return {"ok": False, "error": f"File not found: {path}"}
        ext = target.suffix.lower()
        try:
            import pandas as pd
            df = pd.read_excel(str(target)) if ext in (".xlsx", ".xls", ".xlsm") else pd.read_csv(str(target))
            summary = {
                "rows": int(len(df)),
                "columns": list(map(str, df.columns)),
                "dtypes": {str(k): str(v) for k, v in df.dtypes.items()},
            }
            desc = df.describe(include="all").to_dict()
            # JSON-safe
            clean = {}
            for col, stats in desc.items():
                clean[str(col)] = {str(k): (None if pd.isna(v) else (float(v) if isinstance(v, (int, float)) else str(v)))
                                   for k, v in stats.items()}
            summary["describe"] = clean
            return {"ok": True, "analysis": summary}
        except ImportError:
            # No pandas — read rows via stdlib CSV or openpyxl (Excel), then
            # compute numeric summaries ourselves. Keeps the installer light.
            rows: list = []
            if ext in (".xlsx", ".xlsm"):
                try:
                    import openpyxl
                except ImportError:
                    return {"ok": False, "error": "Excel analysis needs openpyxl. Run: pip install openpyxl"}
                wb = openpyxl.load_workbook(str(target), read_only=True, data_only=True)
                ws = wb.worksheets[0]
                for r in ws.iter_rows(values_only=True):
                    rows.append(["" if c is None else str(c) for c in r])
            else:
                import csv as _csv
                with open(target, newline="", encoding="utf-8", errors="replace") as f:
                    rows = list(_csv.reader(f))
            if not rows:
                return {"ok": True, "analysis": {"rows": 0, "columns": []}}
            header = rows[0]
            data = rows[1:]
            numeric_cols = {}
            for ci, col in enumerate(header):
                vals = []
                for r in data:
                    if ci < len(r):
                        try:
                            vals.append(float(r[ci]))
                        except ValueError:
                            pass
                if vals:
                    numeric_cols[col] = {"min": min(vals), "max": max(vals),
                                         "sum": sum(vals), "mean": sum(vals) / len(vals), "count": len(vals)}
            return {"ok": True, "analysis": {"rows": len(data), "columns": header, "numeric_summary": numeric_cols}}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": f"Could not analyze data: {e}"}


# ── Vision Tool ───────────────────────────────────────────────────────────────

def describe_image(workspace: Path, path: str, question: str = "",
                   context: Optional[dict] = None) -> dict[str, Any]:
    """Understand an image: describe it or answer a question about it.
    Sends the image to the vision model. Degrades gracefully if the active model
    is text-only."""
    if not context or "client" not in context or "model" not in context:
        return {"ok": False, "error": "Vision unavailable (no model context)."}
    try:
        import base64
        target = _safe_path(workspace, path)
        if not target.exists():
            return {"ok": False, "error": f"Image not found: {path}"}
        ext = target.suffix.lower().lstrip(".") or "png"
        mime = "jpeg" if ext in ("jpg", "jpeg") else ext
        b64 = base64.b64encode(target.read_bytes()).decode()
        prompt = question or "Describe this image in detail."
        resp = context["client"].chat.completions.create(
            model=context["model"],
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/{mime};base64,{b64}"}},
                ],
            }],
            max_tokens=1024,
        )
        answer = resp.choices[0].message.content or ""
        return {"ok": True, "description": answer[:8000]}
    except SandboxViolation as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        return {"ok": False, "error": f"Vision failed (the current model may be text-only): {e}"}


# ── Agent Delegation Tool ─────────────────────────────────────────────────────

def ask_user(workspace: Path, question: str, context: Optional[dict] = None) -> dict[str, Any]:
    """Ask the user a clarifying question mid-task and wait for their answer.
    Use when the task is ambiguous and guessing would waste work."""
    ask = context.get("ask_fn") if context else None
    if ask is None:
        return {"ok": False, "error": "Cannot ask the user right now — make a reasonable assumption and continue."}
    if not question.strip():
        return {"ok": False, "error": "question is empty."}
    answer = ask(question, "question", None)
    if not answer:
        return {"ok": True, "answer": "(no answer — proceed with your best judgement)"}
    return {"ok": True, "answer": answer}


def spawn_subagent(workspace: Path, goal: str, context: Optional[dict] = None) -> dict[str, Any]:
    """Delegate a focused subtask to a fresh sub-agent that shares this workspace.
    The sub-agent plans and executes independently, then reports back a summary."""
    if not context or "client" not in context or "model" not in context:
        return {"ok": False, "error": "spawn_subagent unavailable (no model context)."}
    if not goal.strip():
        return {"ok": False, "error": "goal is empty."}
    # Depth guard — a sub-agent may not recursively spawn endless sub-agents.
    depth = int(context.get("depth", 0))
    if depth >= 2:
        return {"ok": False, "error": "Max delegation depth reached — do this subtask yourself."}
    try:
        from agent.loop import run_solo  # lazy import to avoid circular dependency
        sub_caps = dict(context.get("caps", {}))
        sub_caps["max_steps"] = min(sub_caps.get("max_steps", 40), 40)  # keep subtasks bounded
        sub_caps["_subagent_depth"] = depth + 1
        result = run_solo(
            goal, context["model"], workspace, context["client"], sub_caps,
            emit_identity=True, memory=None, skills=None,
        )
        return {"ok": result.get("ok", False), "summary": result.get("summary", "")[:600]}
    except Exception as e:
        return {"ok": False, "error": f"Subagent failed: {e}"}


# ── Tool Schemas ──────────────────────────────────────────────────────────────

def _fn(name, desc, props, required):
    return {"type": "function", "function": {
        "name": name, "description": desc,
        "parameters": {"type": "object", "properties": props, "required": required},
    }}


TOOL_DEFINITIONS = [
    _fn("read_file", "Read a file (up to 16K chars). Use offset/limit for large files.",
        {"path": {"type": "string"}, "offset": {"type": "integer"}, "limit": {"type": "integer"}}, ["path"]),
    _fn("write_file", "Write or overwrite a whole file. Creates parent dirs.",
        {"path": {"type": "string"}, "content": {"type": "string"}}, ["path", "content"]),
    _fn("edit_file", "Precise edit: replace an exact unique substring with new text. Preferred for changing existing files (don't rewrite the whole file).",
        {"path": {"type": "string"}, "old_string": {"type": "string", "description": "Exact text to replace (must be unique unless replace_all)"},
         "new_string": {"type": "string"}, "replace_all": {"type": "boolean"}}, ["path", "old_string", "new_string"]),
    _fn("multi_edit", "Apply several precise edits to one file atomically (all or nothing).",
        {"path": {"type": "string"},
         "edits": {"type": "array", "description": "List of {old_string, new_string, replace_all?}",
                   "items": {"type": "object"}}}, ["path", "edits"]),
    _fn("append_file", "Append text to the end of a file.",
        {"path": {"type": "string"}, "content": {"type": "string"}}, ["path", "content"]),
    _fn("delete_file", "Delete a file or directory.", {"path": {"type": "string"}}, ["path"]),
    _fn("move_file", "Move or rename a file/directory.",
        {"src": {"type": "string"}, "dst": {"type": "string"}}, ["src", "dst"]),
    _fn("copy_file", "Copy a file/directory.",
        {"src": {"type": "string"}, "dst": {"type": "string"}}, ["src", "dst"]),
    _fn("make_dir", "Create a directory (and parents).", {"path": {"type": "string"}}, ["path"]),
    _fn("list_dir", "List files and folders in a directory.", {"path": {"type": "string"}}, []),
    _fn("search_text", "Search file contents by regex (grep). Returns matching lines with file:line.",
        {"pattern": {"type": "string"}, "path": {"type": "string"}, "glob": {"type": "string", "description": "Optional filename filter e.g. '*.py'"},
         "max_results": {"type": "integer"}}, ["pattern"]),
    _fn("find_files", "Find files by glob pattern e.g. '**/*.ts'.",
        {"pattern": {"type": "string"}, "path": {"type": "string"}}, ["pattern"]),
    _fn("shell_exec", "Run a shell command in the workspace (120s, 12K output cap). pip/npm/git/etc.",
        {"command": {"type": "string"}, "timeout": {"type": "integer"}}, ["command"]),
    _fn("run_python", "Run a Python snippet directly and return its output.",
        {"code": {"type": "string"}, "timeout": {"type": "integer"}}, ["code"]),
    _fn("web_fetch", "Fetch a URL (REST APIs, JSON, static HTML; up to 40K chars). Not for JS pages — use browser_action.",
        {"url": {"type": "string"}, "method": {"type": "string", "enum": ["GET", "POST", "PUT", "DELETE", "PATCH"]},
         "headers": {"type": "object"}, "body": {"type": "string"}}, ["url"]),
    _fn("web_search", "Search the web and get a list of result titles + URLs.",
        {"query": {"type": "string"}, "max_results": {"type": "integer"}}, ["query"]),
    _fn("browser_action", "Drive a real headless Chromium browser (JS pages, logins, forms, screenshots).",
        {"action": {"type": "string", "enum": ["navigate", "click", "fill", "get_text", "screenshot", "evaluate", "get_page"]},
         "url": {"type": "string"}, "selector": {"type": "string"}, "value": {"type": "string"},
         "script": {"type": "string"}, "screenshot_path": {"type": "string"}}, ["action"]),
    _fn("generate_image", "Generate an image from a text prompt and save it as a PNG in the workspace.",
        {"prompt": {"type": "string"}, "path": {"type": "string"}, "width": {"type": "integer"}, "height": {"type": "integer"}}, ["prompt"]),
    _fn("read_document", "Read a PDF, Word (.docx), Excel (.xlsx), CSV, or text file and return its text content.",
        {"path": {"type": "string"}}, ["path"]),
    _fn("generate_document", "Create a PDF or Word (.docx) document from text/markdown content. Format chosen by file extension (.pdf or .docx).",
        {"path": {"type": "string"}, "content": {"type": "string"}, "title": {"type": "string"}}, ["path", "content"]),
    _fn("analyze_data", "Analyze a CSV/Excel file: row count, columns, and numeric summaries (min/max/mean/sum).",
        {"path": {"type": "string"}, "question": {"type": "string"}}, ["path"]),
    _fn("describe_image", "Look at an image and describe it or answer a question about it (vision).",
        {"path": {"type": "string"}, "question": {"type": "string"}}, ["path"]),
    _fn("ask_user", "Ask the user a clarifying question mid-task and wait for their answer. Use only when the task is genuinely ambiguous and guessing would waste work.",
        {"question": {"type": "string"}}, ["question"]),
    _fn("spawn_subagent", "Delegate a focused subtask to a fresh sub-agent that shares this workspace and reports back. Use to parallelize or isolate a self-contained chunk of work.",
        {"goal": {"type": "string", "description": "Clear, self-contained subtask description"}}, ["goal"]),
]


# ── Dispatcher ────────────────────────────────────────────────────────────────

def dispatch_tool(workspace: Path, name: str, args: dict, context: Optional[dict] = None) -> dict:
    # Checkpoint: back up anything we're about to change/delete so the run is undoable.
    ckpt = context.get("checkpoint") if context else None

    def _bk(path: str) -> None:
        if ckpt and path:
            try:
                ckpt.backup(_safe_path(workspace, path))
            except Exception:
                pass

    try:
        if name == "read_file":
            return read_file(workspace, args["path"], args.get("offset", 0), args.get("limit", 0))
        if name == "write_file":
            _bk(args["path"])
            return write_file(workspace, args["path"], args.get("content", ""))
        if name == "edit_file":
            _bk(args["path"])
            return edit_file(workspace, args["path"], args["old_string"], args.get("new_string", ""), args.get("replace_all", False))
        if name == "multi_edit":
            _bk(args["path"])
            return multi_edit(workspace, args["path"], args.get("edits", []))
        if name == "append_file":
            _bk(args["path"])
            return append_file(workspace, args["path"], args.get("content", ""))
        if name == "delete_file":
            _bk(args["path"])
            return delete_file(workspace, args["path"])
        if name == "move_file":
            _bk(args["src"]); _bk(args["dst"])
            return move_file(workspace, args["src"], args["dst"])
        if name == "copy_file":
            _bk(args["dst"])
            return copy_file(workspace, args["src"], args["dst"])
        if name == "make_dir":
            return make_dir(workspace, args["path"])
        if name == "list_dir":
            return list_dir(workspace, args.get("path", "."))
        if name == "search_text":
            return search_text(workspace, args["pattern"], args.get("path", "."), args.get("glob", ""), int(args.get("max_results", 100)))
        if name == "find_files":
            return find_files(workspace, args["pattern"], args.get("path", "."))
        if name == "shell_exec":
            return shell_exec(workspace, args["command"], min(int(args.get("timeout", 120)), 300))
        if name == "run_python":
            return run_python(workspace, args["code"], min(int(args.get("timeout", 120)), 300))
        if name == "web_fetch":
            return web_fetch(workspace, args["url"], args.get("method", "GET"), args.get("headers"), args.get("body"))
        if name == "web_search":
            return web_search(workspace, args["query"], int(args.get("max_results", 8)))
        if name == "browser_action":
            return browser_action(workspace, args["action"], args.get("url", ""), args.get("selector", ""),
                                  args.get("value", ""), args.get("script", ""), args.get("screenshot_path", ""))
        if name == "generate_image":
            return generate_image(workspace, args["prompt"], args.get("path", "image.png"),
                                  int(args.get("width", 1024)), int(args.get("height", 1024)))
        if name == "read_document":
            return read_document(workspace, args["path"])
        if name == "generate_document":
            _bk(args["path"])
            return generate_document(workspace, args["path"], args.get("content", ""), args.get("title", ""))
        if name == "analyze_data":
            return analyze_data(workspace, args["path"], args.get("question", ""))
        if name == "describe_image":
            return describe_image(workspace, args["path"], args.get("question", ""), context)
        if name == "ask_user":
            return ask_user(workspace, args["question"], context)
        if name == "spawn_subagent":
            return spawn_subagent(workspace, args["goal"], context)
        return {"ok": False, "error": f"Unknown tool: {name}"}
    except KeyError as e:
        return {"ok": False, "error": f"Missing required argument: {e}"}
    except Exception as e:
        return {"ok": False, "error": f"Tool dispatch error: {e}"}
